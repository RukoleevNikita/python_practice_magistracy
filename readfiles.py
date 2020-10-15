# filesRead = open('F:\программирование\python/read.txt', encoding='utf-8')
# try:
#     text = filesRead.readline()
#     while text:
#         print(text)
#         text = filesRead.readline()
# finally:
#     filesRead.close()

# with open('F:\программирование\python/read.txt', encoding='utf-8') as files:
#     text = files.readline()
#     while text:
#         print(text)
#         text = files.readline()

# модуль python-docx 
import docx
doc = docx.Document('F:\программирование/книги.docx')
# print(doc.paragraphs[0].text) #текст первого абзаца
text = []
for paragraph  in doc.paragraphs:
    text.append(paragraph.text)
    print('\n'.join(text)) # получаем весь текст
